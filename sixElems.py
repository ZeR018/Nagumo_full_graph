import numpy as np
import math

import mainFunks as m
import time
import docx
import joblib
import matplotlib.pyplot as plt

Nstreams = 10
tMax = 1000
highAccuracy = False
k_systems = 4

# Пути для сохранения данных
#path_Doc = './Data/results/' + '4elems_after_G_critical_full_sync_3' + '.docx'
path_Doc = './Data/results/' + 'protazhka_IC_3' + '.docx'

R_data_path = './Data/r_data_protivofaza2.txt'
Graphic_data_path = './Data/graphics/fig'
pathIC_full = './Data/graphics/savedIC_0.02.png'
pathIC = './Data/graphics/IC'
path_LS = './Data/graphics/finalState'
FHN_tr_path = './Data/FHW_coords.txt'

IC_A = np.array([
    1.538879577105834, -2.502158188845857, 0.01, 0.0,
    -2.502158188845857, 2.7191901965415948, 0.01, 0.0,
    2.7191901965415948, 1.724635452347087, 0.01, 0.0,
    1.724635452347087, 1.543752019655198, 0.01, 0.0,
])

#################################################### Functions #################################################

def find_22_mode (G_inh, tMax, highAccuracy = False, doNeedShow = False):
    counter = 0
    while True:
        random_IC = m.IC_random_generator(-3, 3)
        R1_arr, R2_arr, IC = m. \
            make_investigation_of_the_dependence_of_the_order_parameters_on_the_initial_conditions \
            (G_inh, random_IC, tMax, highAccuracy, doNeedShow=doNeedShow)

        eps = 0.1
        if R1_arr[-1] < eps and R2_arr[-1] > 1 - eps:
            return counter, IC, R1_arr, R2_arr
        counter += 1

def find_22_mode_new (G_inh, tMax, highAccuracy = False, doNeedShow=False):
    counter = 0

    while True:
        random_IC = m.IC_FHN_random_generator(FHN_tr_path)
        R1_arr, R2_arr, IC = m.make_investigation_of_the_dependence_of_the_order_parameters_on_the_initial_conditions \
            (G_inh, random_IC, tMax, highAccuracy, doNeedShow=doNeedShow)

        eps = 0.1
        if R1_arr[-1] < eps and R2_arr[-1] > 1 - eps:
            return counter, IC, R1_arr, R2_arr
        counter += 1

# we had 6 elems system => will have 7_elems system with IC (a, b)
def add_elem_to_IC (IC, a, b):
    new_IC = []
    # Копируем старый массив
    for i in range(len(IC)):
        new_IC.append(IC[i])

    new_IC.append(a)
    new_IC.append(b)
    new_IC.append(m.z1_IC)
    new_IC.append(m.z2_IC)

    m.k_systems += 1
    return np.array(new_IC)


def remove_elem_from_IC(IC, num_elem=0):
    if num_elem > m.k_systems:
        print('ur idiot, i cant remove elem. num_elem > k_systems')
        return 0
    else:
        new_IC = []

        for i in range(m.k_systems):
            if i != num_elem:
                for k in range(0, 4):
                    new_IC.append(IC[i * m.k + k])

        m.k_systems -= 1
        return np.array(new_IC)


def existance_investigation_with_changed_IC(index, doNeedShow=False):
    global Nstreams, tMax, highAccuracy

    path_x_start = Graphic_data_path + '_x' + str(index) + '.png'
    path_x_end = Graphic_data_path + '_x_end' + str(index) + '.png'
    path_R = Graphic_data_path + '_R' + str(index) + '.png'
    path_IC = pathIC + str(index) + '.png'
    path_last_state = path_LS + str(index) + '.png'

    G_inh = round(0.24 + 0.001 * index, 6)
    # При маленьких значениях параметра связи берем большое время интегрирования
    if G_inh < 0.005:
        tMax = 10000
    elif G_inh < 0.02:
        tMax = 5000
    else:
        tMax = 500

    last_elem = 149 - 60 + 19
    left_elems = 0
    right_elems = 339
    protivofaza_ic = [left_elems, left_elems,
               right_elems, right_elems]

    full_sync_is = [left_elems - 10, left_elems, left_elems + 10, left_elems + 20]

    IC = m.generate_your_IC_FHN(full_sync_is, pathIC=path_IC, doNeedShow=doNeedShow)
    #highAccuracy = True
    #IC = m.IC_FHN_random_generator(FHN_tr_path, pathSave=path_IC)

    print('Exp: ' + str(index) + '    G_inh: ' + str(G_inh))

    R1_arr, R2_arr, IC, depressed_elements, last_state = m.make_investigation_of_the_dependence_of_the_order_parameters_on_the_initial_conditions \
        (G_inh, IC, tMax, highAccuracy, path_x_start, path_x_end, path_R, path_last_state, doNeedShow)

    return R1_arr, R2_arr, IC, path_x_start, path_x_end, path_R, path_IC, path_last_state, G_inh


def existance_find_cyclop(index, doNeedShow=False):
    global Nstreams, tMax, highAccuracy
    print('Exp: ' + str(index))
    # Костыль
    m.k_systems = 4

    path_x_start = Graphic_data_path + '_x' + str(index) + '.png'
    path_x_end = Graphic_data_path + '_x_end' + str(index) + '.png'
    path_R = Graphic_data_path + '_R' + str(index) + '.png'
    path_IC = pathIC + str(index) + '.png'
    path_last_state = path_LS + str(index) + '.png'

    G_inh = 0.05
    # При маленьких значениях параметра связи берем большое время интегрирования
    if G_inh < 0.005:
        tMax = 10000
    elif G_inh < 0.02:
        tMax = 5000
    else:
        tMax = 1000

    #ind_arr = [1, 2, 3, 339, 340, 341, 169]
    # last_elem = 149 - 60 + 19
    # left_elems = 0 - index
    # right_elems = 339 + index
    # ind_arr = [left_elems - 1, left_elems, left_elems + 1,
    #            right_elems - 1, right_elems, right_elems + 1,
    #            last_elem]
    #
    # IC = m.generate_your_IC_FHN(ind_arr, pathIC=path_IC, doNeedShow=doNeedShow)

    IC = m.IC_FHN_random_generator(FHN_tr_path, pathSave=path_IC)

    R1_arr, R2_arr, IC, depressed_elements, last_state = m.make_investigation_of_the_dependence_of_the_order_parameters_on_the_initial_conditions \
        (G_inh, IC, tMax, highAccuracy, path_x_start, path_x_end, path_R, path_last_state, doNeedShow)

    return R1_arr, R2_arr, IC, path_x_start, path_x_end, path_R, path_IC, path_last_state, G_inh

def existance_show_xt(index, doNeedShow=False):
    global Nstreams, tMax, highAccuracy
    highAccuracy = True

    path_x_start = Graphic_data_path + '_x' + str(index) + '.png'
    path_x_end = Graphic_data_path + '_x_end' + str(index) + '.png'
    path_IC = pathIC + str(index) + '.png'
    path_last_state = path_LS + str(index) + '.png'

    G_inh = round(0.2 + 0.005 * index, 6)
    # При маленьких значениях параметра связи берем большое время интегрирования
    if G_inh < 0.005:
        tMax = 10000
    elif G_inh < 0.02:
        tMax = 5000
    else:
        tMax = 500

    last_elem = 149 - 60 + 19
    left_elems = 0
    right_elems = 339
    protivofaza_ic = [left_elems, left_elems - 10,
                      right_elems, right_elems + 10]

    full_sync_ic = [left_elems - 2, left_elems - 1, left_elems, left_elems + 1]

    IC = m.generate_your_IC_FHN(full_sync_ic, pathIC=path_IC, doNeedShow=doNeedShow)
    # highAccuracy = True
    # IC = m.IC_FHN_random_generator(FHN_tr_path, pathSave=path_IC)

    print('Exp: ' + str(index) + '    G_inh: ' + str(G_inh))

    depressed_elements, last_state = m.make_go_and_show_x_graphics\
        (G_inh, IC, tMax, highAccuracy, path_x_start, path_x_end, path_last_state, doNeedShow)

    return IC, path_x_start, path_x_end, path_IC, path_last_state, G_inh


############################################### Program ####################################################

def make_show_xt():
    maxCount = 50
    m.k_systems = 4

    # Инициализируем файл doc
    mydoc = docx.Document()
    for i in range(0, maxCount):
        existance = joblib.Parallel(n_jobs=Nstreams)(joblib.delayed(existance_show_xt)
                                                     (index)
                                                     for index in range(i * Nstreams, i * Nstreams + Nstreams))

        for j in range(len(existance)):
            IC_i = existance[j][0]
            path_x_start = existance[j][1]
            path_x_end = existance[j][2]
            path_IC = existance[j][3]
            path_last_state = existance[j][4]
            G_inh = existance[j][5]

            # Запись в файл .docx
            mydoc.add_heading('Exp = ' + str(i * Nstreams + j) + ', G_inh = ' + str(G_inh), 2)
            for j in range(0, m.k_systems):
                mydoc.add_paragraph(str(IC_i[j * m.k]) + ', ' + str(IC_i[j * m.k + 1]) + ', ' +
                                    str(IC_i[j * m.k + 2]) + ', ' + str(IC_i[j * m.k + 3]) + ',')

            mydoc.add_picture(path_IC, width=docx.shared.Inches(4))
            mydoc.add_picture(path_x_start, width=docx.shared.Inches(7))
            mydoc.add_picture(path_x_end, width=docx.shared.Inches(7))
            # mydoc.add_picture(path_last_state, width=docx.shared.Inches(5))
            mydoc.add_page_break()
            mydoc.save(path_Doc)

    mydoc.save(path_Doc)

def make_5_7_walk_to_Ginh():

    # Для проверки работы изменения НУ
    #IC_22 = IC_A
    # m.showInitialConditions(IC_22)

    # Убрать элемент
    #IC_22_size_5 = remove_elem_from_IC(IC_22, 2)
    #m.showInitialConditions(IC_22_size_5)

    # Добавить элемент
    # IC_22_size_7 = add_elem_to_IC(IC_22, 1.0, 1.0)
    # print(m.k_systems)
    # m.showInitialConditions(IC_22_size_7)

    # Находим какие-то новые НУ с режимом 2-2
    counter, IC_22, R1_arr_22, R2_arr_22 = find_22_mode(0.02, 500, doNeedShow= False)
    print('mode 2-2 found on the ' + str(counter) + ' attempt')
    m.showInitialConditions(IC_22)

    # G_inh changing params
    class G_inh_cp:
        start = 0.0
        stop = 0.09
        step = 0.01
        num = int((stop - start) / step)

    # Инициализируем файл doc
    mydoc = docx.Document()

    R1_arr = []
    R2_arr = []

    # Changing IC
    #IC_22_size_5 = remove_elem_from_IC(IC_22)
    IC_22_size_7 = add_elem_to_IC(IC_22, 1, -1)

    G_inh_arr = np.linspace(G_inh_cp.start, G_inh_cp.stop, G_inh_cp.num)
    for i in range(0, 5):
        loop_index = i * Nstreams + 1
        existance = joblib.Parallel(n_jobs=Nstreams)(joblib.delayed(existance_investigation_with_changed_IC)
                        (IC_22_size_7, index / 1000.0)
                        for index in range(loop_index, loop_index + Nstreams))

        for j in range(len(existance)):
            R1_arr_i = existance[j][0]
            R2_arr_i = existance[j][1]
            IC = IC_22_size_7
            path_x = existance[j][2]
            path_R = existance[j][3]
            G_inh = existance[j][4]

            # Запись в файл .docx
            mydoc.add_heading('G_inh = ' + str(G_inh))
            mydoc.add_picture(path_x, width=docx.shared.Inches(8))
            mydoc.add_picture(path_R, width=docx.shared.Inches(5))
            mydoc.save(path_Doc)


def make_show_last_state_UC(existance_func):
    maxCount = 30
    m.k_systems = 4

    protivofaza_counter = 0
    full_sync_counter = 0
    other_counter_indexes = []

    # Инициализируем файл doc
    mydoc = docx.Document()
    for i in range(0, maxCount):
        existance = joblib.Parallel(n_jobs=Nstreams)(joblib.delayed(existance_func)
                        (index)
                        for index in range(i * Nstreams, i * Nstreams + Nstreams))

        for j in range(len(existance)):
            # R1_arr, R2_arr, path_x, path_R, path_IC, path_last_state, G_inh
            R1_arr_i = existance[j][0]
            R2_arr_i = existance[j][1]
            IC_i = existance[j][2]
            path_x_start = existance[j][3]
            path_x_end = existance[j][4]
            path_R = existance[j][5]
            path_IC = existance[j][6]
            path_last_state = existance[j][7]
            G_inh = existance[j][8]



            # Запись в файл .docx
            mydoc.add_heading('Exp = ' + str(i * Nstreams + j) + ', G_inh = ' + str(G_inh), 2)
            for j in range(0, m.k_systems):
                mydoc.add_paragraph(str(IC_i[j * m.k]) + ', ' + str(IC_i[j * m.k + 1]) + ', ' +
                                    str(IC_i[j * m.k + 2]) + ', ' + str(IC_i[j * m.k + 3]) + ',')

            mydoc.add_picture(path_IC, width=docx.shared.Inches(4))
            mydoc.add_picture(path_x_start, width=docx.shared.Inches(7))
            mydoc.add_picture(path_x_end, width=docx.shared.Inches(7))
            #mydoc.add_picture(path_last_state, width=docx.shared.Inches(5))
            mydoc.add_picture(path_R, width=docx.shared.Inches(5))
            mydoc.add_page_break()
            mydoc.save(path_Doc)

            eps = 0.03
            if R1_arr_i[-1] >= 1 - eps and R2_arr_i[-1] >= 1 - eps:
                full_sync_counter += 1
            elif R1_arr_i[-1] <= eps and R2_arr_i[-1] >= 1 - eps:
                protivofaza_counter += 1
            else:
                other_counter_indexes.append(i*Nstreams + j)

        print('experiments:', (i+1) * Nstreams, 'full_sync:', full_sync_counter, 'protivofaza:', protivofaza_counter)
        print('unknown:', other_counter_indexes)

    mydoc.add_heading('Results')
    mydoc.add_paragraph('Num experiments:' + str(maxCount*Nstreams))
    mydoc.add_paragraph('Protivofaza:' + str(protivofaza_counter))
    mydoc.add_paragraph('Full_sync:' + str(full_sync_counter))
    for i in range(len(other_counter_indexes)):
        mydoc.add_paragraph('Not Full Sync' + str(other_counter_indexes[i]))
    mydoc.add_page_break()

    mydoc.save(path_Doc)

def make_go_on_last_state():
    last_state = []

    # Определяем НУ для первой итерации
    last_elem = 149 - 60 + 19
    left_elems = 0
    right_elems = 339
    protivofaza_ic = [left_elems, left_elems - 10,
                      right_elems, right_elems + 10]
    full_sync_ic = [left_elems - 2, left_elems - 1, left_elems, left_elems + 1]
    path_IC_0 = pathIC + str(0) + '.png'
    IC = m.generate_your_IC_FHN(full_sync_ic, pathIC=path_IC_0)

    # Инициализируем файл doc
    mydoc = docx.Document()
    # Записываем НУ
    mydoc.add_paragraph('Начальные условия')
    for j in range(0, m.k_systems):
        mydoc.add_paragraph(str(IC[j * m.k]) + ', ' + str(IC[j * m.k + 1]) + ', ' +
                            str(IC[j * m.k + 2]) + ', ' + str(IC[j * m.k + 3]) + ',')
    mydoc.add_picture(path_IC_0)
    mydoc.add_page_break()
    mydoc.save(path_Doc)

    # Сохраняет, в каком состоянии оказались элементы на шаге
    res_states = []
    changing = []
    for i in range(1, 600):
        G_inh_i = 3.0 - i / 50.0

        if G_inh_i <= 0:
            break

        print('Exp ' + str(i) + ', G_inh = ' + str(G_inh_i))

        # При маленьких значениях параметра связи берем большое время интегрирования
        if G_inh_i < 0.005:
            tMax = 5000
        elif G_inh_i < 0.02:
            tMax = 2000
        else:
            tMax = 500

        # Генерируем пути для картинок
        path_x_start = Graphic_data_path + '_x' + str(G_inh_i) + '.png'
        path_x_end = Graphic_data_path + '_x_end' + str(G_inh_i) + '.png'
        path_IC = pathIC + str(G_inh_i) + '.png'
        path_last_state = path_LS + str(G_inh_i) + '.png'

        # Основная функция
        depressed_elements, last_state = m.make_go_and_show_x_graphics(G_inh_i, IC, tMax, highAccuracy_=False,
            path_graph_x_start=path_x_start, path_graph_x_end=path_x_end, path_graph_last_state=path_last_state,
                                                                       doNeedShow=False)
        # Анализируем кол-во синхронизированых между собой элементов
        eps = 0.1
        same_elems_counter = 0
        for i in range(0, k_systems):
            for j in range(0, k_systems):
                if abs(last_state[i*m.k] - last_state[j*m.k]) < eps and i != j:
                    same_elems_counter += 1

        res_states.append(same_elems_counter)
        if len(res_states) > 1:
            if res_states[-1] != res_states[-2]:
                print('on this step stap state was changed G_inh =' + str(G_inh_i))
                print(res_states[-2], res_states[-1])
                changing.append(i)

        # Записываем всё в файл
        mydoc.add_heading('Exp = ' + str(i) + ', G_inh = ' + str(G_inh_i), 2)
        for j in range(0, m.k_systems):
            mydoc.add_paragraph(str(IC[j * m.k]) + ', ' + str(IC[j * m.k + 1]) + ', ' +
                                str(IC[j * m.k + 2]) + ', ' + str(IC[j * m.k + 3]) + ',')
        mydoc.add_picture(path_x_start, width=docx.shared.Inches(6.5))
        mydoc.add_picture(path_x_end, width=docx.shared.Inches(6.5))
        mydoc.add_picture(path_last_state, width=docx.shared.Inches(3))
        mydoc.add_page_break()
        mydoc.save(path_Doc)

        # Берем НУ как последнее состояние предыдущего элемента
        IC = last_state

    if len(changing) != 0:
        mydoc.add_heading('Шаги, на которых происходило изменение состояния:')
        for i in range(len(changing)):
            mydoc.add_paragraph(changing[i])
        mydoc.save(path_Doc)

#make_go_on_last_state()

# IC = np.array([
#     -0.878048233184636, -0.2036613751814247, 0.01, 0,
#     -0.878048233184636, -0.2036613751814247, 0.01, 0,
#     -0.878048233184636, -0.2137613751814247, 0.01, 0,
#     -0.878048233184636, -0.2137613751814247, 0.01, 0])
#
# for i in range(10):
#     G_inh = i / 2.0
#     m.make_go_and_show_x_graphics(G_inh, IC, 500, path_graph_last_state=path_LS + str(i) + '.png',
#                                   path_graph_x_start=(Graphic_data_path + '_x' + str(i) +'.png'),
#                                   path_graph_x_end=(Graphic_data_path + '_x' + str(i) + '_end.png'), highAccuracy_=True)

# left_elems = 0
# right_elems = 339
# protivofaza_ic = [left_elems, left_elems - 10,
#                   right_elems, right_elems + 10]
# full_sync_ic = [left_elems, left_elems, left_elems, left_elems]
# path_IC_0 = pathIC + str(0) + '.png'
# IC = m.generate_your_IC_FHN(full_sync_ic, pathIC=path_IC_0)
# m.make_go_and_show_x_graphics(5.0, IC, 2000, path_graph_last_state=path_LS + '.png',
#     path_graph_x_start=(Graphic_data_path+'_x.png'), path_graph_x_end=(Graphic_data_path+'_x_end.png'), highAccuracy_=True)

'3-1'
left_elems = 0
right_elems = 339
protivofaza_ic = [left_elems, left_elems - 10,
                  right_elems, right_elems + 10]
path_IC_0 = pathIC + str(0) + '.png'
IC = m.generate_your_IC_FHN(protivofaza_ic, pathIC=path_IC_0)
G_inh = 0.187
nondepressed_elem, last_state = m.make_go_and_show_x_graphics(G_inh_, IC, tMax_, highAccuracy_=False, path_graph_x_start=0, path_graph_x_end=0,
                                path_graph_last_state=0, doNeedShow=False)
