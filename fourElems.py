import mainFunks as m
import time
import docx
import joblib
import matplotlib.pyplot as plt

Nstreams = 6
tMax = 3000
highAccuracy = True

# g_ex = []
# g_inh = []
# G_inh = 0.0
# G_ex = 0.02
# for i in range(0, k_systems):
#     g_ex.append([])
#     g_inh.append([])
#     for j in range(0, k):
#         if j == i:
#             g_inh[i].append(0.0)
#             g_ex[i].append(0.0)
#         else:
#             g_inh[i].append(G_inh)
#             g_ex[i].append(G_ex)

# Пути для сохранения данных
R_data_path = './Data/r_data_protivofaza2.txt'
Graphic_data_path = './Data/GraphicsData/saved_fig'
path_Doc = './Data/Docx/' + 'Protivofaza_exp2' + '.docx'

# Инициализируем
final_time = time.time()
#random_IC = m.IC_random_generator(-3, 3)
R1_arr = []
R2_arr = []
IC_arr = []
index = 0
osc_types = []

R1_arr_last = []
R2_arr_last = []
G_inh_arr = []

# Инициализируем файл doc
mydoc = docx.Document()


# Обертка для легкой работы с параллелем, основные изменения (параметр связи, пути сохранения) можно делать здесь
def ex_Ginh_f(index, random_IC, doNeedShow = False):
    if isinstance(random_IC, int or bool):
        random_IC = m.IC_random_generator(-3, 3)
    print('Experiment ' + str(index))
    path_x = Graphic_data_path + '_x' + str(index) + '.png'
    path_R = Graphic_data_path + '_R' + str(index) + '.png'

    G_inh = - index / 1000.0
    #G_inh = -0.045

    # При маленьких значениях параметра связи берем большое время интегрирования
    if G_inh > - 0.005:
        tMax = 25000
    elif G_inh > - 0.02:
        tMax = 8000
    else:
        tMax = 3000

    R1_arr, R2_arr, IC, osc_type = m.\
        make_investigation_of_the_dependence_of_the_order_parameters_on_the_initial_conditions\
        (G_inh, random_IC, tMax, highAccuracy, path_x, path_R, doNeedShow)

    return R1_arr, R2_arr, IC, osc_type, path_x, path_R, G_inh

for i in range(0, 8):
    loop_index = i * Nstreams + 1
    # 0.135 - крайнее значение ингибиторной связи
    #G_inh = i / 1000.0
    random_IC = m.IC_random_generator(-3, 3)

                                                                             # если нужно рандомить каждую итерацию - 0
    # Параллелит                                                             # если один рандом на весь эксп - Random_IC
    existance = joblib.Parallel(n_jobs=Nstreams)(joblib.delayed(ex_Ginh_f)(k, random_IC, False)
                                                 for k in range(loop_index, loop_index + Nstreams))

    # Разбираемся с результатами выполнения каждого потока
    for j in range(0, Nstreams):
        # Записываем данные с каждого потока
        R1_arr_i = existance[j][0]
        R2_arr_i = existance[j][1]
        IC_i = existance[j][2]
        osc_type_i = existance[j][3]
        path_x = existance[j][4]
        path_R = existance[j][5]
        G_inh_i = existance[j][6]

        R1_arr.append(R1_arr_i)
        R2_arr.append(R2_arr_i)
        IC_arr.append(IC_i)
        osc_types.append(osc_type_i)
        G_inh_arr.append(G_inh_i)

        # Запись в файл. Docx
        mydoc.add_heading('Experiment ' + str(loop_index + j), 1)
        mydoc.add_heading("Initial conditions:", 2)
        for j in range(0, m.k_systems):
            mydoc.add_paragraph(str(IC_i[j * m.k_systems]) + ', ' + str(IC_i[j * m.k_systems + 1]) + ', ' +
                                str(IC_i[j * m.k_systems + 2]) + ', ' + str(IC_i[j * m.k_systems + 3]) + ',')
        mydoc.add_picture(path_x, width=docx.shared.Inches(8))
        mydoc.add_picture(path_R, width=docx.shared.Inches(5))
        mydoc.save(path_Doc)

        # Последние элементы массива R1, R2
        R1_arr_last.append(R1_arr_i[-1])
        R2_arr_last.append(R2_arr_i[-1])

    index += 1

# Делаем график зависимости параметров порядка от силы связи
plt.figure()
plt.plot(G_inh_arr, R1_arr_last, label='R1')
plt.plot(G_inh_arr, R2_arr_last, label='R2')
plt.xlabel('G_inh')
plt.ylabel('R_1, R_2')
plt.title('Зависимость R_1, R_2 от G_inh')
plt.legend()
plt.grid()
plt.show()
Graphic_issledovanie_path = './Data/GraphicsData/saved_fig_R_isl.png'
plt.savefig(Graphic_issledovanie_path)
# Добавляем этот график в док
mydoc.add_heading('Final graph', 1)
mydoc.add_picture(Graphic_issledovanie_path, width=docx.shared.Inches(5))
mydoc.save(path_Doc)



# Save data
#m.recordICAndR(R_data_path, R1_arr, IC_arr, G_inh_arr, index)




print('Process end. Final time: ', time.time() - final_time)

